#!/usr/bin/env python3
from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

DEFAULTS = {
    "term": "2025-2026学年第二学期",
    "course": "金融数据分析实验",
    "location": "经济管理学院实验室",
    "student_name": "姚本旭",
    "student_id": "2023060136",
    "teacher": "李老师",
    "date_range": "2026年03月18日---2026年03月18日",
    "title": "基于股票收益率与波动率特征的金融数据分析实验",
    "purpose": (
        "1. 熟悉金融实验中常见数据处理流程，包括样本选取、收益率计算、描述性统计与波动率分析。\n"
        "2. 掌握使用历史价格数据评估资产风险收益特征的基本方法。\n"
        "3. 能够根据实验结果对样本资产的风险水平、收益稳定性和投资特征进行解释。\n"
        "4. 形成规范的实验报告写作习惯，做到结论清晰、过程完整、分析有依据。"
    ),
    "content": (
        "本实验选取某上市公司近一段时期的收盘价数据作为研究样本，围绕收益率与波动率两个核心指标展开分析。\n"
        "首先，对原始价格数据进行整理，剔除缺失值与异常记录，并按照时间顺序计算日收益率。通过收益率序列得到样本的平均收益率、最大值、最小值和标准差等描述性统计指标。\n"
        "其次，结合收益率波动情况观察样本资产在不同交易日中的风险变化特征。实验中发现，样本收益率在多数日期围绕均值上下波动，整体呈现出“平均收益较低但短期波动明显”的特征，这说明资产价格容易受到市场情绪与外部信息冲击。\n"
        "再次，依据波动率指标对该资产的风险进行判断。若标准差较高，则说明收益不稳定、风险偏大；若标准差较低，则表明价格波动相对平缓。通过实验结果可知，该样本资产具备一定投资价值，但在短期持有中仍需关注回撤风险。\n"
        "最后，在实验分析基础上，对金融数据处理方法进行归纳：一是收益率指标能够直观反映投资回报水平；二是波动率能够衡量风险暴露程度；三是将描述性统计与图表观察相结合，有助于更全面地理解资产价格行为。"
    ),
    "summary": (
        "通过本次金融实验，我进一步理解了金融数据分析的基本逻辑，即先完成数据清洗，再进行收益与风险测度，最后依据结果做出解释。实验表明，单纯关注收益水平并不足以支持投资判断，还必须结合波动率与极端值情况综合评估风险。\n"
        "本次实验不仅提升了我使用数据分析方法研究金融问题的能力，也让我认识到规范记录实验步骤和清晰表达分析结论的重要性。后续如果进一步加入多资产比较、回归分析或组合优化方法，实验结果将更具解释力和实践价值。"
    ),
}


def set_para_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_after(anchor, text: str):
    para = anchor._parent.add_paragraph(text)
    para.style = anchor.style
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in para.runs:
        run.font.size = Pt(12)
    anchor._p.addnext(para._p)
    return para


def build_doc(template: Path, output: Path, data: dict[str, str]) -> None:
    doc = Document(str(template))

    replacements = {
        "学年学期：                                   ": f"学年学期：{data['term']}",
        "实验课程名称：                               ": f"实验课程名称：{data['course']}",
        "实验地点：                                   ": f"实验地点：{data['location']}",
        "学生姓名：           姚本旭                  ": f"学生姓名：{data['student_name']}",
        "学    号：         2023060136                ": f"学    号：{data['student_id']}",
        "指导教师：                                   ": f"指导教师：{data['teacher']}",
        "实验时间：   2026年月日---2026年月日": f"实验时间：{data['date_range']}",
    }
    for p in doc.paragraphs:
        if p.text in replacements:
            set_para_text(p, replacements[p.text])

    headings = {p.text.strip(): i for i, p in enumerate(doc.paragraphs)}
    section_map = {
        "一、实验题目": data["title"],
        "二、实验目的和要求": data["purpose"],
        "三、实验内容": data["content"],
        "四、实验总结": data["summary"],
    }
    for heading in ["四、实验总结", "三、实验内容", "二、实验目的和要求", "一、实验题目"]:
        anchor = doc.paragraphs[headings[heading]]
        current = anchor
        for line in section_map[heading].split("\n"):
            current = insert_after(current, line)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", default=str(Path(__file__).resolve().parents[1] / "assets" / "default-template.docx"))
    parser.add_argument("--output", required=True)
    parser.add_argument("--title", default=DEFAULTS["title"])
    parser.add_argument("--term", default=DEFAULTS["term"])
    parser.add_argument("--course", default=DEFAULTS["course"])
    parser.add_argument("--location", default=DEFAULTS["location"])
    parser.add_argument("--student-name", default=DEFAULTS["student_name"])
    parser.add_argument("--student-id", default=DEFAULTS["student_id"])
    parser.add_argument("--teacher", default=DEFAULTS["teacher"])
    parser.add_argument("--date-range", default=DEFAULTS["date_range"])
    parser.add_argument("--purpose", default=DEFAULTS["purpose"])
    parser.add_argument("--content", default=DEFAULTS["content"])
    parser.add_argument("--summary", default=DEFAULTS["summary"])
    args = parser.parse_args()
    build_doc(Path(args.template), Path(args.output), {
        "term": args.term,
        "course": args.course,
        "location": args.location,
        "student_name": args.student_name,
        "student_id": args.student_id,
        "teacher": args.teacher,
        "date_range": args.date_range,
        "title": args.title,
        "purpose": args.purpose,
        "content": args.content,
        "summary": args.summary,
    })
    print(args.output)

if __name__ == "__main__":
    main()
