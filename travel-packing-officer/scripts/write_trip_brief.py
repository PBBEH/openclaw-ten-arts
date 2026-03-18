#!/usr/bin/env python3
from __future__ import annotations
import argparse
from docx import Document
from pathlib import Path


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = 'List Bullet'
        p.add_run(item)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--output', required=True)
    ap.add_argument('--title', required=True)
    ap.add_argument('--summary', nargs='*', default=[])
    ap.add_argument('--flights', nargs='*', default=[])
    ap.add_argument('--hotels', nargs='*', default=[])
    ap.add_argument('--weather', nargs='*', default=[])
    ap.add_argument('--must', nargs='*', default=[])
    ap.add_argument('--optional', nargs='*', default=[])
    ap.add_argument('--reminders', nargs='*', default=[])
    args = ap.parse_args()

    doc = Document()
    doc.add_heading(args.title, 0)
    sections = [
        ('行程摘要', args.summary),
        ('航班建议', args.flights),
        ('酒店建议', args.hotels),
        ('天气与穿搭', args.weather),
        ('必带物品', args.must),
        ('可选物品', args.optional),
        ('出发前提醒', args.reminders),
    ]
    for title, items in sections:
        doc.add_heading(title, level=1)
        add_bullets(doc, items or ['无'])
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)

if __name__ == '__main__':
    main()
